# -*- coding: utf-8 -*-
import sys  

reload(sys)  
sys.setdefaultencoding('utf8')
from django.shortcuts import render
from django.core.mail import send_mail
from django.shortcuts import render
from django.contrib.auth.models import User
from django.http import HttpResponse
import subprocess
import re
import pandas as pd
from subprocess import Popen, PIPE
import uuid
from repindexapi import settings

import docx2txt
from random import randint
from repindexapp.models import *
from django.contrib.auth import logout
from django.contrib.auth import authenticate
from django.http import HttpResponseRedirect
from django.shortcuts import redirect
from django.shortcuts import render_to_response
from django.contrib import messages
from django.contrib import auth
from django.contrib.auth import login 
from django.contrib import messages
from django.views.decorators.cache import cache_control
from django.utils.datastructures import MultiValueDictKeyError
from django.views.decorators.csrf import csrf_exempt
from rest_framework import status
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework.decorators import api_view
from rest_framework import permissions
from rest_framework.decorators import  permission_classes
from datetime import datetime
import csv
import sys
import json
import os
import datetime
from uuid import uuid4
from django.http import JsonResponse
from django.core.files.storage import FileSystemStorage 
from django.core import serializers
from django.contrib.auth.decorators import login_required
import requests
import docx
import unicodedata
import collections
import pypandoc
from django.core.mail import EmailMessage
from uuid import uuid4
from bs4 import BeautifulSoup
import urllib
import dropbox
import subprocess
import json
from tika import parser
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from cStringIO import StringIO
from functools import wraps
import jwt
# Create your views here.
base_url = "/var/www/repindexapi/"
path_url="/var/www/html"


def emailTemplate(firstname,username,password):
    url="http://repindex.com:4200/html/knowledge/"
    emailTemplate.HTML="""
 <div>
<p>Hi """+firstname+""",</p><p>Your Mattersmith  Review account has been successfully created.</p><p>Here are the credentials</p><p><b>Username</b> :  """+username+"""</p><p><b>Password</b> : """+password+"""</p><p>Click on button below to login</p> 
<a href="""+url+""" target="_blank"><button style="background-color:green;color:white">Getting Started</button></a><br>
<br><br>Regards,<br><br><div><hr style="color:red;width:40px;float:left"><br>
<div><i style="color:#003746;font-family:arial;">Harecross,Longframlington, Morpeth NE65 8AN</i><br><i style="color:red">T</i> +44 (0)1665 570815<br><a href="http://www.mattersmith.co.uk" style="font-family:arial; font-size: 14px; line-height: 14px; color: #f0503c;">www.mattersmith.co.uk</a></div><br><a href="http://www.mattersmith.co.uk"><img src="http://www.thinkfarm.co.uk/extranet/Mattersmith/email_signature/images/mattersmith-banner.gif" alt="Mattersmith" style="display:block;" border="0"></a><br>
<div style="font-family:arial; font-size: 12px; line-height: 14px; color: #003746;width:657px">             <hr style="color:#003746"></hr><span style="font-weight: bold;">Authorisation</span><br> 
                <p>This email contains proprietary, privileged and confidential information. If you are not a named addressee (or responsible 
for delivery to a named addressee), you must delete this email and any attachment immediately. We would be grateful if 
you would contact the sender to report any such misdelivery.<p>
                <span style="font-weight: bold;">Privacy and Confidentiality</span><br> 
                <p>This email contains proprietary, privileged and confidential information. If you are not a named addressee (or responsible for delivery to a named addressee), you must delete this email and any attachment immediately. We would be grateful if you would contact the sender to report any such misdelivery.</p> 
                <span style="font-weight: bold;">Email Security and Integrity</span><br> 
                <p>Do not open any attachment to this email without first scanning it for viruses, worms or other unwelcome content. We 
regularly scan and disinfect our computer systems but this is no guarantee that attachments to emails are safe. If you 
cannot determine for yourself that any attachment will not harm your computer systems or data, then delete it immediately 
and ask us to provide the content of the attachment by fax or post. We accept no responsibility for loss, harm or damage 
which results from the failure to follow these instructions.</p>
<hr style="color:#003746"></hr>
</div>
</div>
</div>

    """

def jwtValidator(func):
    @wraps(func)
    def decorator(request, *args, **kwargs):
        if request.META.get('HTTP_AUTHORIZATION') is not None:
           try:
              Token=jwt.decode(request.META.get('HTTP_AUTHORIZATION'), 'secret', algorithms=['HS256'])
           except jwt.DecodeError:
                return JsonResponse({'msg' : 'Token verfication failed'},)
           if Token:
              return func(request, *args, **kwargs)
        else:
            return JsonResponse({'msg' : 'Please provide valid auth token'},)
    return decorator

@csrf_exempt 
def clientregsiter(request,userid):
    print "=======login====user"
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    print "body",body
    username = body['username']
    password = body['password']
    firstname = body['firstname']
    lastname = body['lastname']
    email = body['email']
    mobile_no=body['mobile_no']
    dob=body['dob']
    usertype=body['usertype']
    print body['lastname']
    addedUserCount = UserInfo.objects.filter(addedby=userid).count()
    userLimit = UserInfo.objects.get(userid=userid)
    if addedUserCount > int(userLimit.userlimit)-1:
        print "==========in limit over====="
        jsondata={"msg":"Your Limit is over please exceed it!","success" : "false"}
        return JsonResponse(jsondata)
    alredyexits = User.objects.filter(username=body['username'])[:1]
    if alredyexits:
       jsondata={"msg":"Username alredy exits","success":"false"}
    else:
        userobj= User.objects.create_user(username=username,first_name=firstname,password=password,last_name=lastname, email=email)
        jsondata={"success":"true","msg":"Records created Successfully","userobj.username":str(userobj.username)}
        print "user object",userobj
        getClientid=UserInfo(userobject=userobj,addedby=userid,userid=userobj.id,mobile=mobile_no,dob=dob,user_type_id_id=usertype)
        getClientid.save()
        emailTemplate(firstname,username,password)
	emailStatus = EmailMessage('Your Mattersmith  Review account has been successfully created',emailTemplate.HTML ,"ms-legalservices@mattersmith.co.uk", to=[email])
        emailStatus.content_subtype="html"
        emailStatus.send()
    return JsonResponse(jsondata)

@csrf_exempt
@jwtValidator 
def adminregsiter(request):
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    print "body",body
    username = body['username']
    password = body['password']
    firstname = body['firstname']
    lastname = body['lastname']
    email = body['email']
    userlimit=body['mobile_no']
    dob=""
    usertype=body['usertype']
    print body['lastname']
    alredyexits=User.objects.filter(username=body['username'])[:1]
    if alredyexits:
       jsondata={"msg":"admin alredy exits","success":"false"}
    else:
        userobj= User.objects.create_user(username=username,is_superuser=1,first_name=firstname,password=password,last_name=lastname, email=email)
        jsondata={"success":"true","msg":"Records created Successfully","userobj.username":str(userobj.username)}
        print "user object",userobj
        getClientid=UserInfo(userobject=userobj,userid=userobj.id,userlimit=userlimit,dob=dob,user_type_id_id=usertype)
        getClientid.save()
        emailTemplate(firstname,username,password)
	emailStatus = EmailMessage('Your Mattersmith  Review account has been successfully created',emailTemplate.HTML ,"ms-legalservices@mattersmith.co.uk", to=[email])
        emailStatus.content_subtype="html"
        emailStatus.send()
    return JsonResponse(jsondata)
@csrf_exempt
def Login(request):
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    print "body",body
    print body['username']
    user=authenticate(username=body['username'], password=body['password'])
    if user is not None:
        userInfo =  UserInfo.objects.get(userid = user.id)
        if user.is_active:
           encoded_jwt = jwt.encode({'user': user.username}, 'secret', algorithm='HS256')
           if user.is_superuser:
                 print "super user"
                 login(request, user)
                 print "request", request.user
                 rand_token = uuid4()
                 usertype={"user_type":"admin","user_id":user.id,"username":user.username,"auth_token":encoded_jwt,'concept_status': userInfo.concept_status,'stoplist_status':userInfo.stoplist_status}
                 return JsonResponse(usertype)
                       
           else:
                 print "====else part==="
                 if user.is_staff:
                    print "superuser"
                    login(request, user)
                    print "request",request.user
                    rand_token = uuid4()
                    usertype={"user_type":"superuser","user_id":user.id,"username":user.username,"auth_token":encoded_jwt,'concept_status': userInfo.concept_status,'stoplist_status':userInfo.stoplist_status}
                    return JsonResponse(usertype)
                 else:
                      print "normal user"
                      login(request, user)
                      print "request",request.user
                      rand_token = uuid4()
                      usertype={"user_type":"user","user_id":user.id,"username":user.username,"auth_token":encoded_jwt,'concept_status': userInfo.concept_status,'stoplist_status':userInfo.stoplist_status}
                      return JsonResponse(usertype)
        else:
              msg={"error":"Invalid credentials"}
              return JsonResponse(msg)
    else:
         if user is None:
             msg={"error":"Invalid credentials"}
             return JsonResponse(msg)
@csrf_exempt
def getAllClient(request,userid):
    getclient=UserInfo.objects.filter(addedby=userid).values('mobile','userid', 'userobject__username','userobject__first_name','userobject__last_name','userobject__email','userobject__is_active','created_at')
    getdate=UserInfo.objects.all().values('created_at')
    print str(UserInfo.userid)
    print "type",type(getclient)
    dictionaries = [ obj for obj in getclient ]
    created = [ obj["created_at"] for obj in getdate ]
    print created
    return HttpResponse(json.dumps({"data": dictionaries}, sort_keys=True, default=str),content_type='application/json')
@csrf_exempt
@jwtValidator 
def getAlladmin(request):
    getadmin=User.objects.filter(is_staff=0).values('userinfo__mobile','id','first_name','last_name','username','is_superuser','email','is_active','userinfo__created_at','userinfo__userlimit','userinfo__concept_status','userinfo__stoplist_status')
    getdate=UserInfo.objects.all().values('created_at')
    print str(UserInfo.userid)
    print "type",type(getadmin)
    dictionaries = [ obj for obj in getadmin ]
    created = [ obj["created_at"] for obj in getdate ]
    print created
    return HttpResponse(json.dumps({"data": dictionaries}, sort_keys=True, default=str),content_type='application/json')
@csrf_exempt
def updateClientInfo(request,id):
    print "===id===",id
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    print "body",body
    #username = body['username']
    #password = body['password']
    firstname = body['firstname']
    lastname = body['lastname']
    email = body['email']
    mobile_no=body['mobile_no']
    #dob=body['dob']
    usertype=1
    getclient=UserInfo.objects.get(userid=id)
    getclient.mobile=mobile_no
    #getclient.dob=dob
    getclient.save()
    userobj=User.objects.get(id=id)
    #userobj.username=body['username']
    #userobj.set_password(body['password'])
    userobj.first_name=body['firstname']
    userobj.last_name=body['lastname']
    userobj.email=body['email']
    userobj.save()
    msg={"succss":"true","msg":"Record updated succssfully"}
    return JsonResponse(msg)
@csrf_exempt
def updatesuperuser(request,id):
    print "============in updatesuperuser========"
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    print "body",body
    firstname = body['firstname']
    lastname = body['lastname']
    email = body['email']
    userobj=User.objects.get(id=id)
    userobj.first_name=body['firstname']
    userobj.last_name=body['lastname']
    userobj.email=body['email']
    userobj.save()
    msg={"succss":"true","msg":"Record updated succssfully"}
    return JsonResponse(msg)

@csrf_exempt
@jwtValidator
def addProject(request,user_id):
    print "==========userid====",user_id
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    if Projects.objects.filter(name=body['project_name'],created_by_id=user_id).exists():
           msg={"error":"Project Name already exits","success":"false"}
           return JsonResponse(msg)
    project = Projects(status=1)

    project.name = body['project_name']
    project.created_by = project.updated_by = UserInfo.objects.get(userid=user_id)
    project.save()
    msg={"Succss":"Project Name saved Succssfully","success":"true"}
    return JsonResponse(msg)
@csrf_exempt
def getProject(request,user_id):
    allProject=Projects.objects.filter(created_by=user_id)
    jsondata = serializers.serialize("json", allProject)
    return HttpResponse(jsondata,status=200,content_type='application/json')

@csrf_exempt
@jwtValidator
def getprojectDetails(request, project_id):
    project = Projects.objects.get(id = project_id)
    folder_type_list = FolderType.objects.filter(project__id=project_id).values_list('folder__id', flat=True)
    folders = Folder.objects.filter(id__in=folder_type_list)
    text_list_data =  json.loads(serializers.serialize("json", TextList.objects.filter(project_id = project_id)))
    urls = Url.objects.filter(project_id = project_id)
    urls = json.loads(serializers.serialize('json', urls))
    #print "folders>>",folders
    folders_list = []
    temp_dict = {}
    for f in folders:
        temp_dict = {}
        temp_dict['folder'] = f.name
        temp_dict['id'] = f.id
        tl_objs=TextList.objects.filter(folder_id=f)
        tl_tmp=[]
        for t in tl_objs:
            tl_tmp.append({'id':t.id,'file_name':t.name})
        temp_dict['tl'] = tl_tmp
        temp_dict['sub_folders'] = [create_folder_dict(request,Folder.objects.filter(parent_folder_id=f)) if Folder.objects.filter(parent_folder_id=f).exists() else ""]
        folders_list.append(temp_dict)
    jsond = json.dumps({'urls': urls, 'project_id': project_id, 'text_folders': json.loads(serializers.serialize("json", Folder.objects.filter(project_id = project_id))) ,'text_list_data': text_list_data, 'text_list_folders': folders_list,'concept_lists':getConceptLists(request,project_id),'stop_lists':getStopLists(request,3)})
    #msg={"project_id":project_id}
    print jsond
    return HttpResponse(jsond)

def getStopLists(request,project_id):
    lst_stop_lists = []
    try:
        lst_stoplist_obj = StopList.objects.filter(id=project_id)
        for x in lst_stoplist_obj:
            dict_stop_list = {}
            lst_stop_list_values = []
            lst_stoplist_value_obj = StopListValue.objects.filter(stop_list=x)
            for y in lst_stoplist_value_obj:
                lst_stop_list_values.append(y.value)
            dict_stop_list.update({"name": x.name, "id": x.id, "values": lst_stop_list_values})
            lst_stop_lists.append(dict_stop_list)
    except Exception, ex13:
            print  ex13
    return lst_stop_lists

def getConceptLists(request,project_id):
    project = Projects.objects.get(id=project_id)
    try:
        lst_concept_lists = []
        lst_concept_list_obj = ConceptList.objects.filter(project_id=project)
        for x in lst_concept_list_obj:
            dict_concept_list = {}
            lst_concept_list_values = []
            lst_concept_list_value_obj = ConceptListValue.objects.filter(concept_list=x)
            lst_umbrella = []
            for y in lst_concept_list_value_obj:

                dict_umbrella = {}
                dict_umbrella.update({"umbrella": y.umbrella})
                    # lst_umbrella.append(y.umbrella)
                lst_concept_list_umbrella_values_obj = ConceptListUmbrellaValues.objects.filter(concept_list_value=y)
                for z in lst_concept_list_umbrella_values_obj:
                    lst_terms_obj = ConceptListUmbrellaValues.objects.filter(
                        concept_list_value=z.concept_list_value)
                    terms = ""

                    lst_terms=[]
                    for term_obj in lst_terms_obj:

                        terms = terms + term_obj.term + ","
                        lst_terms.append(term_obj.term)
                    terms = terms[:-1]
                    dict_umbrella.update({"terms": terms, "id": z.id,"lst_terms":lst_terms})
                lst_umbrella.append(dict_umbrella)
            dict_concept_list.update({"name": x.name,'folder_id':x.folder_id.id, "values": lst_umbrella, "id": x.id})
            lst_concept_lists.append(dict_concept_list)

    except Exception, ex23:
           print  ex23
    #print "concept_lists=====", json.dumps(lst_concept_lists)
    return lst_concept_lists

def create_folder_dict(request, folders):

            folders_dict = {}
            for f in folders:
                print f
                folders_dict['folder'] = f.name
                folders_dict['id'] = f.id
                tl_objs = TextList.objects.filter(folder_id=f)
                tl_tmp = []
                for t in tl_objs:
                    tl_tmp.append({'id': t.id, 'file_name': t.name})

                folders_dict['tl'] = tl_tmp

                #folders_dict['tl'] = list(TextList.objects.filter(folder_id=f))
                folders_dict['sub_folders'] = [create_folder_dict(Folder.objects.filter(parent_folder_id=f)) if Folder.objects.filter(parent_folder_id=f).exists() else ""]
            return folders_dict
@csrf_exempt
def addFolder(request,project_id,user_id):
    print "======projectid",project_id
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    try:
        if 'folder_id' in body:
            folder_id=body['folder_id']
            if folder_id !="/":
                folder_id=int(folder_id.split("/")[0].strip())
                parent_folder = Folder.objects.get(id=folder_id)
            else:
                    parent_folder = None
        else:
            parent_folder = None

        folder = Folder(status=1)
        folder.name = body['folder_name']
        folder.project_id = project_id
        folder.created_by = folder.updated_by = UserInfo.objects.get(id=user_id)
        folder.status = 1
        folder.save()

        if parent_folder:
            folder.parent_folder_id = parent_folder
            folder.save()
        else:
            print "project id=",project_id
            # map with folder type
            foldertype_obj = FolderType(folder=folder, project_id=project_id,type="text-list")
            foldertype_obj.save()
            print folder, folder.id
        msg={"msg":"Folder created successfully","success":"true"}
        return JsonResponse(msg)
    except Exception,ex:
           print ex
def docx_to_text(request,filename):
        doc = docx.Document(filename)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        return '\n'.join(fullText)


def doc_to_text_catdoc(request):
        (fi, fo, fe) = os.popen3('catdoc -w "%s"' % '/tmp/temp_doc')

        fi.close()
        retval = fo.read()
        erroroutput = fe.read()
        fo.close()
        fe.close()
        return retval
        
        if not erroroutput:
            return retval
        else:
            raise OSError("Executing the command caused an error: %s" % erroroutput)
       
@csrf_exempt
def Savetextlist(request,project_id,user_id,folder_id,access_token):
        parent_folder=None
        #body_unicode = request.body.decode("utf-8")
        #body = json.loads(body_unicode)
        print ">>>>>>>>>>>>>",request.FILES,len(request.FILES)
        pfolder_id=folder_id
        # File upload
        if pfolder_id:
            #parent_folder = Folder.objects.get(id=request.POST['folder_id'])
            parent_folder = Folder.objects.get(id=pfolder_id)
            print "file name in folder id", request.FILES
	mag = {}    
            
        if len(request.FILES) > 0:
                for i in request.FILES:
                    file_type=""
                    try:
                        file_type=request.FILES[i]._name.strip().split(".")[
                            len(request.FILES[i]._name.strip().split("."))-1]
                    except Exception,ex:
                        print ex
                    if 'msword' in request.FILES[i].content_type:
                        file_contents,html = "not found","not found"
                        myfile = request.FILES[i]
                        directory=settings.MEDIA_URL + str(uuid.uuid4()) + "/"
                        fs = FileSystemStorage(location=directory)
                        if not os.path.exists(directory):
                            os.makedirs(directory)
                            filename = fs.save(myfile.name, myfile)
                            uploaded_file_url = fs.url(directory+filename)
                        access_token_value =access_token
                        if not access_token_value=="undefined":
                            file_from = directory+filename
                            file_to="/"+filename
                            dbx = dropbox.Dropbox(access_token_value)
                            f = open(file_from, 'rb')
                            dbx.files_upload(f.read(), file_to)
                        print subprocess.call(['soffice', '--headless', '--convert-to', 'docx', directory+filename])
                        file_contentswithslah=document_to_text(request,filename,directory+filename)
                        print subprocess.call(['soffice', '--headless', '--convert-to', 'docx', directory+filename])
                        file_contents=file_contentswithslah.replace("|"," ")
                        import mammoth
                        with open(base_url+filename+"x", "rb") as docx_file:
                            result = mammoth.convert_to_html(docx_file)
                            html = result.value # The generated HTML
                            messages = result.messages # Any messages, such as warnings during conversion
                        print "------>htmldata ===>=======> ",html
                        with_br_tag=html
                    elif file_type=="docx":
                        myfile     = request.FILES[i]
                        directory  = settings.MEDIA_URL + str(uuid.uuid4()) + "/"
                        fs = FileSystemStorage(location=directory)
                        if not os.path.exists(directory):
                            os.makedirs(directory)
                            filename = fs.save(myfile.name, myfile)
                        uploaded_file_url = fs.url(directory+filename)
                        import mammoth
                        with open(directory+filename, "rb") as docx_file:
                            result = mammoth.convert_to_html(docx_file)
                            html = result.value
                            messages = result.messages
                        file_contents=docx2txt.process(request.FILES[i])
                        with_br_tag=html
                        import dropbox
                        access_token_value =access_token
                        if not access_token_value=="undefined":
                            file_from = directory+filename
                            file_to="/"+filename
                            dbx = dropbox.Dropbox(access_token_value)
                            f = open(file_from, 'rb')
                            dbx.files_upload(f.read(), file_to)
                    elif file_type == "txt":
                        file_contents  = request.FILES[i].file.read()
                        with_br_tag  = file_contents
                        print "------------------",file_contents
		    elif file_type == "pdf":
			 #import pdb
			 #pdb.set_trace()
                         myfile = request.FILES[i]
                         directory  = settings.MEDIA_URL + str(uuid.uuid4()) + "/"
                         fs = FileSystemStorage(location=directory)
                         filename = fs.save(myfile.name, myfile)
                         uploaded_file_url = fs.url(directory+filename)
                         pdfData=document_to_text(request,filename,directory+filename)
                         print "pdf -----content----",pdfData
                         file_contents = pdfData
                         with_br_tag = pdfData
                         #parsedPDF = parser.from_file(directory+filename)
                         #file_contents = parsedPDF["content"]
                         #with_br_tag = parsedPDF["content"]
                    elif file_type == "csv":
                         print "====in csv file======="
                         import csv
                         myfile     = request.FILES[i]
                         directory  = settings.MEDIA_URL + str(uuid.uuid4()) + "/"
                         fs = FileSystemStorage(location=directory)
                         if not os.path.exists(directory):
                            os.makedirs(directory)
                            filename = fs.save(myfile.name, myfile)
                         uploaded_file_url = fs.url(directory+filename)
                         data=[]
                         with open(directory+filename) as File:
                               reader = csv.reader(File)
                               for row in reader:
                                   data.append(row)

                         newvalue=''
                         for i in data:
                              print "value of i",i
                              newvalue+=' '.join(i)

                         file_contents=newvalue
                         with_br_tag=newvalue
                
                
                    text_list_obj = TextList(name=request.FILES[i]._name,
                        folder_id=parent_folder,project_id_id=project_id,status=1)
                    text_list_obj.created_by = text_list_obj.updated_by = UserInfo.objects.get(id=user_id)
                    text_list_obj.save()
                    text_list_detail_obj = TextListDetails(text_list_id=text_list_obj,file_content=with_br_tag,
                        file_content_plain_text=file_contents,status=1)
                    text_list_detail_obj.created_by = text_list_detail_obj.updated_by = UserInfo.objects.get(id=user_id)
                    if  text_list_detail_obj.save():
			 msg={"succss":"Textlist created successfully","success":"true"}
		    else:
			 msg={"succss":"Textlist not created","success":"false"}
		print '------------msg---------',msg
                return JsonResponse(msg)
        else:
            msg={"succss":"Textlist not created","success":"false"}
        return JsonResponse(msg)
def document_to_text(request,filename,file_path):
    
    if filename[-4:] == ".doc":
        cmd = ['antiword', file_path]
        p = Popen(cmd, stdout=PIPE,shell=True)
        stdout, stderr = p.communicate()
        return stdout.decode('ascii', 'ignore')
    elif filename[-4:] == ".pdf":
         return convert_pdf_to_txt(file_path)
"""
@csrf_exempt
def Savetextlist(request,project_id,user_id):
        parent_folder=None
        #body_unicode = request.body.decode("utf-8")
        #body = json.loads(body_unicode)
        
            
        if len(request.FILES) > 0:
            
            print "file save in outside folder id",request.FILES
            if len(request.FILES) >0:
                file_type=""
                try:
                    print "files name====>",request.FILES['text_file']._name.strip().split(".")
                    file_type=request.FILES['text_file']._name.strip().split(".")[
                        len(request.FILES['text_file']._name.strip().split("."))-1]
                except Exception,ex:
                    print ex
                print "============>",file_type
               
                file_contents = request.FILES['text_file'].file.read()
               
                if 'msword' in request.FILES['text_file'].content_type:
                    file_contents="not found"

                    html="not found"
                    myfile = request.FILES['text_file']
                    directory=settings.MEDIA_URL + str(uuid.uuid4()) + "/" 
                    print "diretory========",directory
                    fs = FileSystemStorage(location=directory)
                    if not os.path.exists(directory):
                        os.makedirs(directory)
                        filename = fs.save(myfile.name, myfile)
                        print "filename",filename
                        uploaded_file_url = fs.url(directory+filename)
                    print "uploaded_file_url=====",uploaded_file_url
                    print "actula path",directory+filename
                    file_contentswithslah=document_to_text(request,filename,directory+filename)
                   
                    file_contents=file_contentswithslah.replace("|"," ")
                     #print "===========file_content=======",file_contents
                    html=file_contents.replace("\n","<br>")
                    print html
                elif file_type=="docx":
                     file_obj = open('temp_doc', 'w')
                     file_obj.write(file_contents)
                     file_obj.close()
                     myfile = request.FILES['text_file']
                     directory=settings.MEDIA_URL + str(uuid.uuid4()) + "/"
                     fs = FileSystemStorage(location=directory)
                     if not os.path.exists(directory):
                        os.makedirs(directory)
                        filename = fs.save(myfile.name, myfile)
                        print "filename",filename
                        uploaded_file_url = fs.url(directory+filename)
                        print "uploaded_file_url=====",uploaded_file_url
                        print directory+filename
                        html = convert(directory+filename)
                        file_contents=docx2txt.process(request.FILES['text_file'])
                        print html
                     
                #print "file_contents===========================>",file_contents
                
                text_list_obj = TextList(name=request.FILES['text_file']._name,
                    folder_id=parent_folder,project_id_id=project_id,status=1)
                text_list_obj.created_by = text_list_obj.updated_by = UserInfo.objects.get(id=user_id)
                text_list_obj.save()
                with_br_tag = html
                text_list_detail_obj = TextListDetails(text_list_id=text_list_obj,file_content=with_br_tag,
                    file_content_plain_text=file_contents,status=1)
                #print "with_br_tag======================>",with_br_tag
                text_list_detail_obj.created_by = text_list_detail_obj.updated_by = UserInfo.objects.get(id=user_id)
                text_list_detail_obj.save()
                msg={"succss":"Textlist created successfully","success":"true"}
                return JsonResponse(msg)
        else:
            pass
        msg={"succss":"Textlist created successfully","success":"true"}
        return JsonResponse(msg)
"""
@csrf_exempt
def saveConceptlist(request,project_id,user_id,folder_id):
    #print "9999999999999999", request.POST

        file_name=request.FILES['browse_concept_list']._name
        print "file name=",file_name
        print "concept list>>>>>>>", request.FILES
        if request.FILES:
            file_contents = request.FILES['browse_concept_list'].file.read()
            if file_contents.count('->') > 200:
           
               msg={"success":"alert"}
               return JsonResponse(msg)
            print "file_contents=",file_contents
            lst_concept_list=file_contents.split("\n")
            conceptObject=ConceptList.objects.filter(project_id_id=project_id).values('name','status')
            print "======conceptObject name======",conceptObject
            coceptNames=[concept['name'] for concept in conceptObject]
            print "====list of files",coceptNames
            if file_name in coceptNames:
               print "in =====if========="
               jsondata={"msg":"concept list alredy exits","success":"false"}
               return JsonResponse(jsondata)
            else:
                 concept_list_obj=ConceptList(folder_id_id=folder_id,name=file_name,project_id=Projects.objects.get(id=project_id),status=1,created_on=datetime.datetime.now(),updated_on=datetime.datetime.now())
                 concept_list_obj.save()
                 for x in lst_concept_list:
                
                     if len(x)==0:
                    
                         msg={"succss":"conceptlist added successfully"}
                         return JsonResponse(msg)
                     else: 
                         
                         umbrella=ConceptListValue(concept_list=concept_list_obj,umbrella=x.split("[")[0],status=1)
                         umbrella.save()
                         try:
                            for y in x.split("->")[1].split(","):
                                if len(y.strip())>0 and y is not None:
                                   try:
                                      defination=x[x.find("[")+1 : x.find("]")]
                                      #print "=========defination",defination
                                      #print "============terms=======",y.strip()
                                      umbrella_value_obj=ConceptListUmbrellaValues(concept_list_value=umbrella,term=y.strip(),defination=defination)
                                      umbrella_value_obj.save()
                                   except Exception,ex21:
                                          print "===ex2=====",ex21

                       
                         except IndexError:

                                msg={"succss":"conceptlist added successfully"}
                                JsonResponse(msg)
                

            





        context = {
            'project': Projects.objects.get(id=project_id),
            'text_list_folders': get_side_bar(request,project_id),
        }

        msg={"succss":"conceptlist added successfully"}
        return JsonResponse(msg)
def get_side_bar(request,p_id):
        project = Projects.objects.get(id=p_id)
        folder_type_list = FolderType.objects.filter(project__id=p_id).values_list('folder__id', flat=True)
        folders = Folder.objects.filter(id__in=folder_type_list)
        #print "folders>>",folders
        folders_list = []
        temp_dict = {}
        for f in folders:
            temp_dict = {}
            temp_dict['folder'] = f.name
            temp_dict['id'] = f.id
            tl_objs=TextList.objects.filter(folder_id=f)
            tl_tmp=[]
            for t in tl_objs:
                tl_tmp.append({'id':t.id,'file_name':t.name})

            temp_dict['tl'] = tl_tmp
            temp_dict['sub_folders'] = [create_folder_dict(Folder.objects.filter(request,parent_folder_id=f)) if Folder.objects.filter(parent_folder_id=f).exists() else ""]
            folders_list.append(temp_dict)
            #print "temp_dict>>>>>",temp_dict
        #print "folders_list>>>",folders_list
        return folders_list
def create_folder_dict(request, folders):

            folders_dict = {}
            for f in folders:
                print f
                folders_dict['folder'] = f.name
                folders_dict['id'] = f.id
                tl_objs = TextList.objects.filter(folder_id=f)
                tl_tmp = []
                for t in tl_objs:
                    tl_tmp.append({'id': t.id, 'file_name': t.name})

                folders_dict['tl'] = tl_tmp

                #folders_dict['tl'] = list(TextList.objects.filter(folder_id=f))
                folders_dict['sub_folders'] = [create_folder_dict(Folder.objects.filter(parent_folder_id=f)) if Folder.objects.filter(parent_folder_id=f).exists() else ""]
            return folders_dict

@csrf_exempt

def deleteUser(request):
    "====delete user==="
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    user_id=body["user_id"]
    getclient=User.objects.filter(id=user_id).delete();
    getUserInfo = UserInfo.objects.filter(userid=user_id).delete();
    msg={"msg":"records deleted successfully","user_id":user_id,"succss":"true"}
    return JsonResponse(msg)

@csrf_exempt
def inactiveUser(request,user_id):
    userobj=User.objects.get(id=user_id)
    userobj.is_active=False
    userobj.save()
    msg={"msg":"User Inactive succssfully","user_id":user_id,"succss":"true"}
    return JsonResponse(msg)

@csrf_exempt
def activateUser(request,user_id):
    userobj=User.objects.get(id=user_id)
    userobj.is_active=True
    userobj.save()
    msg={"msg":"User activted succssfully","user_id":user_id,"succss":"true"}
    return JsonResponse(msg)

def getuserbyid(request,user_id):
    getclient=UserInfo.objects.filter(userid=user_id).values('mobile','userid', 'userobject__username','userobject__first_name','userobject__last_name','userobject__email','userobject__is_active','created_at')
    getdate=UserInfo.objects.all().values('created_at')
    
    print str(UserInfo.userid)
   
    print "type",type(getclient)
    dictionaries = [ obj for obj in getclient ]
    created = [ obj["created_at"] for obj in getdate ]
    
    print created
    
    return HttpResponse(json.dumps({"data": dictionaries}, sort_keys=True, default=str),content_type='application/json')

@csrf_exempt    
def concordanceline(request,project_id):
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    selected_files = unicodedata.normalize('NFKD', body['selected_ids']).encode('ascii','ignore').split(",")
    selected_files = filter(None, selected_files)
    exportCSV=[]
    data_list = []
    for file in selected_files:

        file_contents = TextListDetails.objects.get(text_list_id__id=file).file_content_plain_text
        file_contents_data = file_contents
        file_contents_data = file_contents.replace("%",'')
        print "file_contents_data========",file_contents_data
        textlistname=TextList.objects.get(id=file)

        toi_str = unicodedata.normalize('NFKD', body["txt_toi"] ).encode('ascii','ignore')
        import copy
        #termHeadValue=copy.deepcopy(lst_toi_str)  
        lst_toi_str = toi_str.split(',')
        termHeadValue=copy.deepcopy(lst_toi_str)
        # print lst_toi_str,'***********'
        lst_toi_str = list(set([x.lower() for x in lst_toi_str])) #
        termId=ConceptListUmbrellaValues.objects.filter(term=lst_toi_str[0])[0]
        termHead=ConceptListUmbrellaValues.objects.filter(concept_list_value_id=termId.concept_list_value_id)[0]
        print "termid===>",termId.concept_list_value_id
        conceptId= ConceptListValue.objects.get(id=termId.concept_list_value_id)
        print "term head===>",termHead.term
        print "concpetlist id",conceptId.concept_list_id
        allConceptsValues = getConceptListWordsWithHead(request,conceptId.concept_list_id)
        #terms_obj = ConceptListUmbrellaValues.objects.filter(concept_list_value=y)
        Conceptslist=[]
        #print "allConceptsValues==========>",allConceptsValues
        for i in allConceptsValues:
            Conceptslist.append(i['values'])
        print Conceptslist

        for toi_lower in lst_toi_str:
            sidesData  = {}            
            rightData = {}
            allData = []
            line_number=0
            # if toi_lower.count(' ') > 0:                     
            pattern = toi_lower.lower() #
            findData = [(m.start(0), m.end(0)) for m in re.finditer(pattern, file_contents_data.lower())]
            if len(findData):
                for ele in findData:                        
                    left  = ele[0]
                    right = ele[1]
                    frontData = left-1 #
                    if not (file_contents_data[right].isalpha() or file_contents_data[frontData].isalpha()): #
                        totalLeftData = file_contents_data[0:left]
                        totalRightData = file_contents_data[right:]
                        if totalLeftData.count(" ") < 50:
                            sidesData['left_data'] = file_contents_data[0:]
                        else:
                            spaceCount = 0
                            flag = True
                            tempLeft = left
                            while flag:
                                if file_contents_data[tempLeft] == ' ':
                                    spaceCount += 1
                                if spaceCount == 49:
                                    flag = False
                                tempLeft = tempLeft - 1
                            sidesData['left_data'] = file_contents_data[tempLeft:left]
                        if totalRightData.count(" ") < 50:
                            sidesData['right_data'] = file_contents_data[ele[1]:]
                        else:
                            spaceCount = 0
                            flag = True
                            tempRight = right
                            while flag:
                                if file_contents_data[tempRight] == ' ':
                                    spaceCount += 1
                                if spaceCount == 49:
                                    flag = False
                                tempRight = tempRight + 1
                            sidesData['right_data'] = file_contents_data[right:tempRight]
                        data_list.append({'toi':toi_lower.lower(),
                                        'combination_concept':CombinationHead(request,sidesData['left_data'],sidesData['right_data'],Conceptslist),
                                        'left':sidesData['left_data'],
                                        'head':termHeadValue[0].upper(),
                                        'right':sidesData['right_data'],
                                        'line_number':line_number,
                                        "file_id":file,
                                        'text_list':textlistname.name}
                                    )

                        exportCSV.append([
                                            textlistname.name,
                                            ' '.join(sidesData['left_data'].split(" ")[0:10]),
                                            toi_lower,
                                            ' '.join(sidesData['right_data'].split(" ")[0:10]),
                                            lst_toi_str[0].upper(),
                                            ConcordeWith(request,sidesData['left_data'],sidesData['right_data'],lst_toi_str),
                                            ])

                        # exportCSV.append({'Toi':toi_lower.lower(),
                        #                 'Combination Concept':ConcordeWith(request,sidesData['left_data'],sidesData['right_data'],lst_toi_str),
                        #                 'Left':' '.join(sidesData['left_data'].split(" ")[0:10]),
                        #                 'Head':lst_toi_str[0].upper(),
                        #                 'Right':' '.join(sidesData['right_data'].split(" ")[0:10]),
                        #                 'Text List':textlistname.name}
                        #             )

        continue
    context = {
            'results': data_list,
            'result_type': "Concordance Lines",
            'queries': QueryList.objects.filter(project_id_id=project_id),
            'file_id':selected_files[0],
            'totaltoi':toi_str
    }
    col = ['Text List','Left','Search','Right','Head','Combination']#
    df = pd.DataFrame(exportCSV,columns=col)
    df.to_csv(base_url +"ConcordanceResult.csv", index=False,header=True)
    os.system("mv "+base_url +"ConcordanceResult.csv "+ path_url)
    jsond = json.dumps({'totaltoi': toi_str,'head':lst_toi_str[0], 'file_id[0]': selected_files,'results':data_list})
    return HttpResponse(jsond)
        
def getStopLists(request,sl_id):

        lst_stop_lists = []
        try:


            lst_stoplist_obj = StopList.objects.filter(id=sl_id)
            for x in lst_stoplist_obj:

                dict_stop_list = {}
                lst_stop_list_values = []
                lst_stoplist_value_obj = StopListValue.objects.filter(stop_list=x)
                for y in lst_stoplist_value_obj:
                    lst_stop_list_values.append(y.value)

                dict_stop_list.update({"name": x.name, "id": x.id, "values": lst_stop_list_values})
                lst_stop_lists.append(dict_stop_list)
        except Exception, ex13:
            print  ex13
        return lst_stop_lists

@csrf_exempt    
def importstoplist(request):
    file_name=request.FILES['browse_stop_list']
    print "file_name======",file_name
    if  request.FILES:
        file_contents = request.FILES['browse_stop_list'].file.read()
        print "file_contents=",file_contents
        lst_stop_list=file_contents.split("\n")
        try:

            stop_list_obj=StopList(name=file_name,project_id=Projects.objects.get(id=4),status=1,created_on=datetime.datetime.now(),updated_on=datetime.datetime.now())
            stop_list_obj.save()
            for x in lst_stop_list:
                print "=================x============",x
                if len(x)==0:
                    print "=========================in x===",len(x),type(x)
                    msg={"succss":"conceptlist added successfully"}
                    return JsonResponse(msg)
                stop_list_value=StopListValue(stop_list=stop_list_obj,value=x.split("->")[0],status=1)
                stop_list_value.save()
        except Exception,ex78:
               print ex78

        context = {
            'project': Projects.objects.get(id=project_id),
            'text_list_folders': get_side_bar(request,project_id),

        }

    jsond = json.dumps({'project': project, 'text_list_folders': text_list_folders})
    return HttpResponse(jsond)

@csrf_exempt

def renameproject(request):
     print "==============>"
     body_unicode = request.body.decode("utf-8")
     body = json.loads(body_unicode)
     projectId=body['project_id']
     projectName=body['project_name']  
    
     gotid=Projects.objects.get(id=projectId);
     gotid.name = projectName;
     gotid.save();
     msg={"msg":"project rename succssfully","succss":"true"}
     return JsonResponse(msg)

@csrf_exempt
def deleteproject(request):
     body_unicode = request.body.decode("utf-8")
     body = json.loads(body_unicode)
     projectid=body['project_id']
      
     gotid=Projects.objects.get(id=projectid);
     gotid.delete();
     msg={"msg":"delete project succssfully","succss":"true"}
     return JsonResponse(msg)

@csrf_exempt
def renamefolder(request,folder_id):
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    textobj=Folder.objects.get(id=folder_id)
    textobj.name=body["folder_name"]
    textobj.save()
    msg={"msg":"rename project succssfully","succss":"true"}
    return JsonResponse(msg)

@csrf_exempt
def deletefolder(request,folder_id):
   
    textobj=Folder.objects.filter(id=folder_id).delete()
   
    msg={"msg":"rename project succssfully","succss":"true"}
    return JsonResponse(msg)

@csrf_exempt
def renameconcept(request):
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    conceptid=body['concept_id']
    rename=body['rename_list']
    print "=============id============",conceptid,rename
    gotid=ConceptList.objects.get(id=conceptid);
    gotid.name = rename;
    project_id = gotid.project_id.id;
    gotid.save();
    msg={"msg":"rename concept succssfully","succss":"true"}
    return JsonResponse(msg)
    

@csrf_exempt
def deleteconcept(request):
     body_unicode = request.body.decode("utf-8")
     body = json.loads(body_unicode)
     conceptid=body['concept_id']
     print "=============id============",conceptid
     gotid=ConceptList.objects.get(id=conceptid);
     project_id = gotid.project_id.id;
     gotid.delete();
     msg={"msg":"delete concept succssfully","succss":"true"}
     return JsonResponse(msg)

@csrf_exempt
@jwtValidator
def renameproject(request):
     print "==============>"
     body_unicode = request.body.decode("utf-8")
     body = json.loads(body_unicode)
     projectId=body['project_id']
     projectName=body['project_name']  
    
     gotid=Projects.objects.get(id=projectId);
     gotid.name = projectName;
     gotid.save();
     msg={"msg":"project rename succssfully","succss":"true"}
     return JsonResponse(msg)

@csrf_exempt
@jwtValidator
def deleteproject(request):
     body_unicode = request.body.decode("utf-8")
     body = json.loads(body_unicode)
     projectid=body['project_id']
      
     gotid=Projects.objects.get(id=projectid);
     gotid.delete();
     msg={"msg":"delete project succssfully","succss":"true"}
     return JsonResponse(msg)

@csrf_exempt
def selfcordoncordance(request,project_id,user_id):
    #xyz
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    selected_files = unicodedata.normalize('NFKD', body['selected_ids']).encode('ascii','ignore').split(",")
    selected_files = filter(None, selected_files)
    data_list = []
    downloadCsv=[]
    for ids in selected_files:
        file_contents = TextListDetails.objects.get(text_list_id__id=ids).file_content_plain_text
        file_contents_data = unicodedata.normalize('NFKD', file_contents).encode('ascii','ignore')
        file_contents_data = file_contents.replace("%",'')
        textlistname=TextList.objects.get(id=ids)
        cl_id = body['hf_cmb_concept_lists']
        allConceptsValues = getConceptListWordsWithHead(request,cl_id)
        print "allConceptsValues",allConceptsValues
        #terms_obj = ConceptListUmbrellaValues.objects.filter(concept_list_value=y)
        Conceptslist=[]
        #print "allConceptsValues==========>",allConceptsValues
        for i in allConceptsValues:
            Conceptslist.append(i['values'])
        print Conceptslist
        lst_toi_str = getConceptListWordsWithHead(request,cl_id)
        file_content_length = len(file_contents)
        for data in lst_toi_str:
            line_number = 0
            # toi_str = unicodedata.normalize('NFKD', data["values"]).encode('ascii','ignore').split(",")
            toi_str = data["values"].encode('ascii','ignore').split(",")
            toi_str = filter(None, toi_str)
            
            toi_str_lower = list(set([l.lower() for l in toi_str])) #
            for toi_lower in toi_str_lower:
                sidesData  = {}
                
                rightData = {}
                allData = []
                # if toi_lower.count(' ') > 0:                     
                pattern = toi_lower #
                findData = [(m.start(0), m.end(0)) for m in re.finditer(pattern, file_contents_data.lower())]
                if len(findData):
                    for ele in findData:                        
                        left  = ele[0]
                        right = ele[1]
                        frontData = left-1 #
                        
                        if not (file_contents_data[right].isalpha() or file_contents_data[frontData].isalpha() ): #
                            totalLeftData = file_contents_data[0:left]
                            totalRightData = file_contents_data[right:]
                            if totalLeftData.count(" ") < 50:
                                sidesData['left_data'] = file_contents_data[0:]
                            else:
                                spaceCount = 0
                                flag = True
                                tempLeft = left
                                while flag:
                                    if file_contents_data[tempLeft] == ' ':
                                        spaceCount += 1
                                    if spaceCount == 49:
                                        flag = False
                                    tempLeft = tempLeft - 1
                                sidesData['left_data'] = file_contents_data[tempLeft:left]
                            if totalRightData.count(" ") < 50:
                                sidesData['right_data'] = file_contents_data[ele[1]:]
                            else:
                                spaceCount = 0
                                flag = True
                                tempRight = right
                                while flag:
                                    if file_contents_data[tempRight] == ' ':
                                        spaceCount += 1
                                    if spaceCount == 49:
                                        flag = False
                                    tempRight = tempRight + 1
                                sidesData['right_data'] = file_contents_data[right:tempRight]
                            data_list.append({'toi':toi_lower,
                                            'combination_concept':CombinationHead(request,sidesData['left_data'],sidesData['right_data'],Conceptslist),
                                            'left':sidesData['left_data'],
                                            'head':data["head"].upper(),
                                            'right':sidesData['right_data'],
                                            'line_number':line_number,
                                            "file_id":ids,
                                            'text_list':textlistname.name}
                                        )
                            downloadCsv.append([
                                                textlistname.name,
                                                ' '.join(sidesData['left_data'].split(" ")[0:10]),
                                                toi_lower,
                                                ' '.join(sidesData['right_data'].split(" ")[0:10]),
                                                data["head"].upper(),
                                                checkConcordeWith(request,sidesData['left_data'],sidesData['right_data'],lst_toi_str),
                                                ])
                            
                            # downloadCsv.append({'TOI':toi_lower,
                            #                 'Combination concept':checkConcordeWith(request,sidesData['left_data'],sidesData['right_data'],lst_toi_str),
                            #                 'Left':' '.join(sidesData['left_data'].split(" ")[0:10]),
                            #                 'Head':data["head"].upper(),
                            #                 'Right':' '.join(sidesData['right_data'].split(" ")[0:10]),
                            #                 'Text List':textlistname.name}
                            #             )

            continue    
                
    context = {
            'results': data_list,
            'result_type': "self Concordance Lines",
            'project': Projects.objects.get(id=project_id),           
            'file_id': selected_files
    }
    col = ['Text List','Left','Search','Right','Head','Combination']#
    df = pd.DataFrame(downloadCsv,columns=col)
    df.to_csv(base_url +'exportResult.csv', index=False,header=True)
    os.system("mv "+base_url +"exportResult.csv ")
    jsond = json.dumps({'totaltoi': lst_toi_str, 'file_id': selected_files,'results':data_list})
    return HttpResponse(jsond)

def checkConcordeWith(request,left,right,concepts):
        pass
        try:

            for q in concepts:
                toi_str = q["values"]
                toi_str_org = toi_str
                lstToi = toi_str.split(" ")
                for z in lstToi:
                    if z in map(lambda name: name.upper(), left.split(" ")):
                        return q["head"].upper()
                    elif z.upper() in map(lambda name: name.upper(), right.split(" ")):
                        return q["head"].upper()
        except Exception,ex:
            print ex
        return ""

def ConcordeWith(request,left,right,concepts):
        pass
        try:
                lstToi = concepts.split(",")
                for z in lstToi:
                    if z in  left.split(" "):
                       return z.upper()
                    elif z in right.split(" "):
                       return z.upper()
        except Exception,ex:
            print ex
        return ""

def getConceptListWordsWithHead(request, cl_id):
        clwh = []
        try:
            cl_obj = ConceptList.objects.get(id=cl_id)
            umbrella = ConceptListValue.objects.filter(concept_list=cl_obj)
            for y in umbrella:
                terms_obj = ConceptListUmbrellaValues.objects.filter(concept_list_value=y)
                concepts = ""
                str = ""
                for z in terms_obj:
                    str = str + z.term + "," 
                clwh.append({"head":y.umbrella,"values":str})
            return clwh
        except Exception,ex2:
            print ex2

@csrf_exempt
def clusterquery(request,project_id,user_id):
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    selected_files = body['selected_ids'].split(',')[:-1]
    print selected_files
    data_list = []
    for file in selected_files:

        file_contents = TextListDetails.objects.get(text_list_id__id=file).file_content_plain_text
        textlistname=TextList.objects.get(id=file)
        fc=file_contents.encode('utf-8')
        query_string=body['txtarea_cluster_pairs']
        lst_dimensions = re.findall(r'(\w+):', query_string)
        lst_attributes = re.findall(r':(\w+)', query_string)
        lst_concepts = re.findall(r"(?<={)[^}]*(?=})", query_string)

        print "dimensions", lst_dimensions
        print "attributes=", lst_attributes
        print "concepts=", lst_concepts
        dimension_dict = {}
        lst_cluster_splits = []
        count = 0
        dimension = ""
        s1 = query_string.split(":")
        print "s1=", s1
        print "LEN OF s1=", len(s1)
        for x in s1:
            count = count + 1
            dimension_dict = {}
            if count != len(s1):
                print "cnt=", count
                if count == 1:
                    dimension = x
                    attribute_dict = {}
                    tmp_lst_concepts = re.findall(r"(?<={)[^}]*(?=})", s1[count])
                    attribute = s1[count].split('{')[0]
                    concepts = []
                    for y in tmp_lst_concepts:
                        lst_concepts = y.split("<>")
                        for z in lst_concepts:
                            if len(z) != 0:
                                concepts.append(z)
                    attribute_dict.update({"attribute": attribute, "concepts": concepts})

                    print "dimension_dict==", dimension_dict
                    dimension_dict.update({"dimension": dimension, "attributes": attribute_dict})

                    lst_cluster_splits.append(dimension_dict)

                    print "########", s1[count].split('}')

                    if count != len(s1) - 1:
                        dimension = attribute = s1[count].split('}')[len(s1[count].split('}')) - 1]
                        print ">>>>>>>>>>>dimension=", dimension

            else:
                print "cnt=", count
                attribute = x.split('{')[0]
                concepts = []
                attribute_dict = {}
                tmp_lst_concepts = re.findall(r"(?<={)[^}]*(?=})", x)
                for y in tmp_lst_concepts:
                    lst_concepts = y.split("<>")
                    for z in lst_concepts:
                        if len(z) != 0:
                            concepts.append(z)

                attribute_dict.update({"attribute": attribute, "concepts": concepts})

        for clsr in lst_cluster_splits:
            toi_str=""

            for cn in clsr['attributes']['concepts']:
                toi_str=toi_str+cn+" "

            listText=fc.split(".")



            cat_str = request.POST.get('txt_cat')

            toi_str_org=toi_str

            lstToi=toi_str.split(" ")

            nw=7
            cl=[]
            line_number=0

            for y in lstToi:

                if len(y.strip())==0:
                    continue


                if(y.upper()=="AND"):
                    pass

                line_number=0
                newstr=y.lower()
                print "====newstr word=======>",newstr
                print "=====file_contents=======",file_contents
                file_contentlower=file_contents.lower()
                cc=newstr.split(" ")
                print "==============ccc===========",len(cc)
                
                if newstr in file_contentlower:
                   print "oooooogot search words=============>",newstr,type(newstr)
                   indexgot=[x.start() for x in re.finditer(newstr, file_contentlower)]
                   for valueoffirst in indexgot:
                         leftvalue=valueoffirst-500
                         rightvalue=valueoffirst+500
                         left=file_contents[leftvalue:valueoffirst]
                         line_number=randint(1, 40)
                         right=file_contents[valueoffirst+len(newstr):rightvalue]
                         data_list.append({'toi':newstr,'left':left,'right':right,'line_number':line_number,"file_id":file,'text_list':textlistname.name})
                    
            continue;    
                
        context = {
            'results': data_list,
            'result_type': "Cluster query list",
            'project': Projects.objects.get(id=project_id),
            'file_id': selected_files[0]
        }
    jsond = json.dumps({ 'file_id': selected_files,'results':data_list})
    return HttpResponse(jsond)

@csrf_exempt
@jwtValidator
def getalltextlist(request,project_id):  
    textlist=TextList.objects.filter(project_id=project_id)
    jsondata = serializers.serialize("json", textlist)
    return HttpResponse(jsondata,status=200,content_type='application/json')
    
@csrf_exempt
def getconceptlist(request,project_id):
    print "====project_id===",project_id
    conceptlist=ConceptList.objects.filter(project_id=project_id)
    jsondata = serializers.serialize("json", conceptlist)
    return HttpResponse(jsondata,status=200,content_type='application/json')

@csrf_exempt
def getfilecontes(request,file_id):
    data=TextListDetails.objects.filter(text_list_id=file_id)
    jsondata = serializers.serialize("json", data)
    return HttpResponse(jsondata,status=200,content_type='application/json')

@csrf_exempt
def frequencycount(request):
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    requestData = dict([(str(k), str(v)) for k, v in body.items()])
    rowDocIds = requestData['selected_ids'].split(',')
    docIds = filter(None, rowDocIds)
    file_contents = ''
    for ids in docIds:
        file_contents = file_contents + (TextListDetails.objects.get(text_list_id=ids).file_content_plain_text + " ").encode('ascii','ignore')
    # import pdb
    # pdb.set_trace()
    if  requestData['stoplist_id'] != 'None':
        stopListId   = requestData['stoplist_id']
        stopListData = StopListValue.objects.filter(stop_list=stopListId).values('value')
        stopValue    = [value.values()[0].encode('ascii','ignore').lower() for value in stopListData]
    
    #get umbrella values
    conceptListId = requestData['cmb_concept_lists']
    allUmbrell = ConceptListValue.objects.filter(concept_list=conceptListId).values('id','umbrella')
    umbrellData = []
    zeroUmbrellaCountData = []
    tempFile = file_contents.lower()
    tempUmbrellaTermName = []
    columnDataCsv = []
    for data in allUmbrell:
        tId, umbrellaName = data['id'], data['umbrella'].encode('ascii','ignore')
        tempUmbrellaTermName.append(umbrellaName.lower())
        termData = ConceptListUmbrellaValues.objects.filter(concept_list_value=tId).values('term')
        umbrellaCount = {}
        countData =0
        listContent = [] #
        for term in termData:
            termName =  term['term'].encode('ascii','ignore').lower()
            if not termName in listContent: #
                if not termName in tempUmbrellaTermName: #
                    tempUmbrellaTermName.append(termName)
                findData = [(m.start(0), m.end(0)) for m in re.finditer(termName, tempFile)]
                for data in findData:
                    frontData = data[0] - 1 #
                    if not ((ord(tempFile[data[1]]) >= 97 and ord(tempFile[data[1]]) <= 122) or (ord(tempFile[frontData]) >= 97 and ord(tempFile[frontData]) <=122)) :#
                        countData = countData + 1
            listContent.append(termName)#
        if countData > 0:
            columnDataCsv.append([umbrellaName,countData])
            umbrellData.append({'count': countData,'term': umbrellaName })
        else:
            columnDataCsv.append([umbrellaName,countData])
            zeroUmbrellaCountData.append({'count': countData,'term': umbrellaName })
    umbrellData.extend(zeroUmbrellaCountData)
    # umbrellData.append({'term': '-----------Other Words----------','count': '----------' })

    # for stop in stopValue:
    #     if ' '+stop+' ' in tempFile:
    #         tempFile = tempFile.replace(' '+stop+' ',' ')
    # commonWordData = []
    # for i in re.finditer(r'\w+', tempFile):
    #     data = i.group(0)
    #     commonWordCount = 0
    #     if not data in tempUmbrellaTermName:
    #         commonWordCount=len([(m.start(0), m.end(0)) for m in re.finditer(' '+termName, tempFile)])
    #     commonWordData.append({'term': data,'count': commonWordCount })
    # umbrellData.extend(commonWordData)

    df = pd.DataFrame.from_dict(umbrellData)
    df = pd.DataFrame(columnDataCsv,columns=['Term','Count'])
    df.to_csv(base_url +'FrequencyResult.csv', index=False,header=True)
    os.system("mv "+base_url +"FrequencyResult.csv "+path_url)

    jsond = json.dumps({ 'results': umbrellData})
    return HttpResponse(jsond)

    
@csrf_exempt
def getstoplistdata(request):
    print "=========in get stop list*********8==========="
    stoplist=StopList.objects.all()
    jsondata = serializers.serialize("json", stoplist)
    return HttpResponse(jsondata,status=200,content_type='application/json')

@csrf_exempt
def savedocumentdata(request):
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    fileid=body["fileid"]
    data=body["data"]
    t = TextListDetails.objects.get(text_list_id=fileid)
    t.file_content =data   # change field
    t.save()
    msg={"msg":"Document updated succssfully","succss":"true"}
    return JsonResponse(msg)

@csrf_exempt
def renameStopList(request):    
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    stopListId=body['id']
    newName=body['name']    
    gotid=StopList.objects.get(id=stopListId);
    gotid.name = newName;
    gotid.save();
    msg={"msg":"rename stoplist succssfully","succss":"true"}
    return JsonResponse(msg)

@csrf_exempt
def getConceptListWords(request,cl_id):
        try:
            pass
            cl_obj=ConceptList.objects.get(id=cl_id)
            strData=""
            data = {}
            umbrella=ConceptListValue.objects.filter(concept_list=cl_obj)
            
            for y in umbrella:
                listData = []
                terms_obj=ConceptListUmbrellaValues.objects.filter(concept_list_value=y)
                for z in terms_obj:
                    listData.append(unicodedata.normalize('NFKD', z.term).encode('ascii','ignore'))
                    strData=strData+z.term + ","
                data[listData[0]] = listData
            return  strData,data
        except Exception,ex3:
            print ex3

@csrf_exempt
def deleteStopList(request, stoplist_id):   
    gotid=StopList.objects.get(id=int(stoplist_id));
    gotid.delete();
    msg={"msg":"stoplist delete succssfully","succss":"true"}
    return JsonResponse(msg)

@csrf_exempt
def drop(request):    
    print "request url===========>", request.get_full_path()
    url=request.get_full_path()
    auth_code=url.split("=")[1]
    print "===auth_code",auth_code
    direct_output = subprocess.check_output("curl https://api.dropbox.com/oauth2/token -d code="+auth_code+" -d grant_type=authorization_code -d redirect_uri=http://localhost:8000/repindexapi/drop -u 4l6wn6uf43jv5mq:nno61ecrv636xci", shell=True)
    print "====>direct output===>",direct_output
    data = json.loads(direct_output)
    #print "=========>",data["access_token"]


    """
    access_token = 'Ce_-cjcQyGAAAAAAAAAATx-vDds4dFmMtYX4xj7ocy0W7_QcVWUvmkU3G9J-xlxl'
    file_from = '/home/aniket/temp.doc' 
    file_to="/ccc.txt"
    dbx = dropbox.Dropbox(access_token)
    f = open(file_from, 'rb')
    dbx.files_upload(f.read(), file_to)

    print "response ",dbx
    """
    from django.shortcuts import redirect
    params={"access_token":"xnakssnjsdn","url":"http://localhost:4200/#/project"}
    # return JsonResponse(params)
    return redirect('http://localhost:4200/#/project/access_token/'+data["access_token"])
    #return HttpResponseRedirect("http://localhost:4200/#/project",{'code':"abcxyz"})

@csrf_exempt
def dropboxapi(request):
    import webbrowser
    access_token="9IZM-WIUMdAAAAAAAAAACbKR-ADFiSfEoM_bYxBIz8tNAF4dKEx6Y26J7xbg-V4I"
    file_from = '/home/heaptrace/demo123.docx' 
    file_to="/ddc.txt"
    dbx = dropbox.Dropbox(access_token)
    f = open(file_from, 'rb')
    print "----",f
    dbx.files_upload(f.read(), file_to)
    #client = dropbox.client.DropboxClient("Ce_-cjcQyGAAAAAAAAAADlQtBrr4Sp5uEftLvpnioWua_haYXIos2gws5Z1izq9A")
    #api key=jcg41kabcs0rx4a
    #scret= krueas6but8ifjp
    #HxHXsvLyvxAAAAAAAAAAQ4zdRcaz3dHknpksp9JkoJwK2XWah3pzT_Hp_XOBasa6
    #webbrowser.open("https://www.dropbox.com/oauth2/authorize?client_id=4l6wn6uf43jv5mq&response_type=code&redirect_uri=http://localhost:8000/repindexapi/drop")

    #curl https://api.dropbox.com/oauth2/token -d code=Ce_-cjcQyGAAAAAAAAAANYLccbODOHr1GWizbp8jl-M -d grant_type=authorization_code -d redirect_uri=http://localhost:8000/repindexapi/drop -u 4l6wn6uf43jv5mq:nno61ecrv636xci
    #Ce_-cjcQyGAAAAAAAAAANOUqmEmBQBlnS2zOCLucWuk
    
def wordexport(request):
    from StringIO import StringIO    
    import pycurl

    url = 'http://www.google.com/'

    storage = StringIO()
    c = pycurl.Curl()
    c.setopt(c.URL, url)
    c.setopt(c.WRITEFUNCTION, storage.write)
    c.perform()
    c.close()
    content = storage.getvalue()
    print content
@csrf_exempt
def addDefination(request):
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    id=body['id']
    defi=body['def']
    print "====id===",id
    print "====defi===",defi
    conceptUmbrall=ConceptListUmbrellaValues.objects.get(id=id)
    conceptUmbrall.defination=defi
    conceptUmbrall.save()
    msg={"msg":"defination added succssfully","succss":"true"}
    return JsonResponse(msg)
@csrf_exempt
def getDefination(request,id):
    defination=ConceptListUmbrellaValues.objects.filter(id=id)
    print "defination========>",defination
    jsondata = serializers.serialize("json", defination)
    return HttpResponse(jsondata,status=200,content_type='application/json')
@csrf_exempt
def exportdoc(request):
     xlsx_data="<h2>hello  wordl</h2>"
     print "ddddddddddddddddd",xlsx_data
     response = HttpResponse(xlsx_data, content_type='application/msword')
     #response['Content-Length'] = os.path.getsize(filename)
     return response

@csrf_exempt
def downloadDocument(request,file_id):
    settings_dir = os.path.dirname(__file__)
    PROJECT_ROOT = os.path.abspath(os.path.dirname(settings_dir))
    ids =  str(file_id)
    file_contents = TextListDetails.objects.get(text_list_id=ids).file_content
    # print file_contents
    inputPath  = PROJECT_ROOT+'/static/downloads/temp.html'
    outputPath = PROJECT_ROOT+'/static/downloads/output.docx'
    if os.path.exists(outputPath):
        os.unlink(outputPath)
    if not  os.path.exists(PROJECT_ROOT+'/static'):
        os.makedirs(PROJECT_ROOT+'/static/downloads')
        os.chmod(PROJECT_ROOT+'/static', 0o777)
        os.chmod(PROJECT_ROOT+'/static', 0o777)
    if not  os.path.exists(PROJECT_ROOT+'/static/downloads'):
        os.makedirs(PROJECT_ROOT+'/static/downloads')
        os.chmod(PROJECT_ROOT+'/static', 0o777)
        os.chmod(PROJECT_ROOT+'/static/downloads', 0o777)
    if not os.path.exists(inputPath):
        try:
            file = open(inputPath,'a')
            file.close()
        except:
            print "no file created"
    fw = open(inputPath,'w')
    fw.write(file_contents.encode('utf-8'))
    fw.close()    
    output = pypandoc.convert(source=inputPath,
                             format='html',
                             to='docx',
                             outputfile=outputPath,
                             extra_args=['-RTS']
                             )
    from django.contrib.staticfiles.templatetags.staticfiles import static
    filename = outputPath
    data = open(filename, "rb").read()
    response = HttpResponse( data, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')     
    response['Content-Disposition'] = 'attachment; filename="demo.docx"'
    return response
 
@csrf_exempt 
def forgotpassword(request):
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode) 
    print "========",body["email"]
    email=body['email']
    exits = User.objects.filter(email=body['email'])[:1]
    print "================>",exits
    if exits:
       print "** Records exits **"
       rand_token = uuid4()
       print rand_token
       exitsid = User.objects.get(email=body['email'])
       print "id===================>",exitsid.id
       UserInfoobj=UserInfo.objects.get(userid=exitsid.id)
       print "userobject",UserInfoobj
       UserInfoobj.token=rand_token
       UserInfoobj.save()
       print"===1===="
       email = EmailMessage('Reset your Repindex account password', "Please click on the below link for reset password<\n https://asconceptreviewer.com/#/login/token/"+str(rand_token),"ms-legalservices@mattersmith.co.uk" ,to=[email])
       if email.send():
          print"===2===="
          msg={"msg":" Email send succssfully !! Please check in your inbox","succss":"true"}
          return JsonResponse(msg)
    else:
       msg={"msg":"Email with this user not associated","succss":"false"}
       return JsonResponse(msg)
    return JsonResponse(msg)

@csrf_exempt 
def changepassword(request,token):
    print "===========",token
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    print "=========",body["password"] 
    UserInfoobj=UserInfo.objects.get(token=token)
    print "==============",UserInfoobj
    userobj=User.objects.get(id=UserInfoobj.id)
    userobj.set_password(body["password"] )
    userobj.save()
    msg={"msg":"Password Updated succssfully","succss":"true"}
    return JsonResponse(msg)

@csrf_exempt
def deleteTextList(request, textlist_id):
    gotid=TextList.objects.get(id=int(textlist_id));
    gotid.delete()
    msg={"msg":"stoplist delete succssfully","succss":"true"}
    return JsonResponse(msg)

@csrf_exempt    
def importmasterlist(request,user_id):
        file_name=request.FILES['text_file']._name
        if request.FILES:
            file_contents = request.FILES['text_file'].file.read()
            print file_contents
            #import pdb
            #pdb.set_trace()
            lst_concept_list=file_contents.split("\n")
            print lst_concept_list
            project_id = Projects.objects.values('id')[0]['id']
            print project_id
            concept_list_obj= ConceptList(name = file_name,project_id=Projects.objects.get(id=project_id),status=0,master_list_status=1,created_on=datetime.datetime.now(),updated_on=datetime.datetime.now())
            concept_list_obj.save()
            for x in lst_concept_list:                
                if len(x)==0:                    
                    msg={"succss":"conceptlist added successfully"}
                    return JsonResponse(msg)
                else: 
                    umbrella=ConceptListValue(concept_list=concept_list_obj,umbrella=x.split("[")[0],status=1)
                    umbrella.save()
                    try:
                       for y in x.split("->")[1].split(","):
                           if len(y.strip())>0 and y is not None:
                              try:
                                  defination=x[x.find("[")+1 : x.find("]")]
                                  umbrella_value_obj=ConceptListUmbrellaValues(concept_list_value=umbrella,term=y.strip(),defination=defination)
                                  umbrella_value_obj.save()
                              except Exception,ex21:
                                        print "===ex2=====",ex21
                    except IndexError:
                            msg={"succss":"conceptlist added successfully"}
                            JsonResponse(msg)
            context = {
            'project': Projects.objects.get(id=project_id),
            'text_list_folders': get_side_bar(request,project_id),
             }

            msg={"succss":"conceptlist added successfully"}
            return JsonResponse(msg)
            
@csrf_exempt    
def getmasterlist(request):
    data = ConceptList.objects.filter(master_list_status=1)
    ids  = data.values('id')
    project_id = data[0].project_id_id
    allData = getNewConceptLists(request, project_id)
    idList = []
    allInfo = []
    for i in ids:
        idList.append(i['id'])   
    for x in allData:
        if x['id'] in idList:
            allInfo.append(x)
    msg = {"result":'Data not found'}
    if len(data) > 0:
        # jsondata = serializers.serialize("json", allInfo)
        return HttpResponse(json.dumps(allInfo) ,status = 200, content_type = 'application/json')
    else:
        return JsonResponse(msg)
        
@csrf_exempt
def deleteMasterList(request, id):
    print "====>",id
    recordId = ConceptList.objects.get(id = int(id));
    recordId.delete()
    msg={"msg":"masterlist delete surepindexapp_conceptlistccssfully","succss":"true"}
    return JsonResponse(msg)
    
@csrf_exempt
def updateMasterList(request):
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)    
    objectData  =  ConceptList.objects.get(id=int(body['id']))    
    objectData.name = body['name']
    objectData.save()
    msg = {"msg" : "masterlist rename successfully", "success":True}
    return JsonResponse(msg)
 
@csrf_exempt
def deleteMultipleTextlist(request):
    print "hello";
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)  
    print body["filesid"].split(",")
    textIds=body["filesid"].split(",")
    for i in textIds:
         ids=TextList.objects.get(id=int(i));
         ids.delete()
    msg = {"msg" : "delete files successfully", "success":True}
    return JsonResponse(msg) 
@csrf_exempt
def deleteMultipleConceptlist(request):
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    conceptid=body['concept_id'].split(",")
    for i in conceptid:
        print "=============id============",i
        cids=ConceptList.objects.get(id=i);
        project_id = cids.project_id.id;
        cids.delete();
    msg={"msg":"delete concept succssfully","succss":"true"}
    return JsonResponse(msg)

@csrf_exempt
def getValue(request):
     body_unicode = request.body.decode("utf-8")
     body = json.loads(body_unicode)
     filename=body["filename"].split(',')
     leftandright=' '.join(body["leftandright"].split(' ')[2:6])
     print "==============>",leftandright
     file_id=body['file_id']
     file_contents = TextListDetails.objects.get(text_list_id__id=file_id).file_content_plain_text
     indexposition=file_contents.find(leftandright)
     file_content=file_contents[indexposition:indexposition+500]
     pound=u'\xa3'
     file_content=file_content.replace(pound,"$")
     filesplit=file_content.split(" ")
     values=re.sub("\$([1-9][0-9,]+)", "",file_content)
     modiflyvalues=values.split(" ")
     loopvalues=[i for i in filesplit if i not in modiflyvalues]
     print "loopvalues---checking-->",loopvalues
     if len(loopvalues) > 0:
        print "=============>in loopvalues",loopvalues
     else:
        print "===========esle part"

        findingDoller=body["leftandright"].split(' ')
        CurrenyNames=["dollar","dollars","pound","pounds","amount of","USD","SGD","EUR","euros","CNY","CAN","NZD","JPY","yen","RAND","CHF","franc","TWD","NT","AUD","SEK","Krona","rupee","rupees","contract","contract of","value","value of","values at","valued at","valuing at","valuation","valuation of","range","range of","ranges from","ranged from","ranging from","rent","rent of","rental","rental of"]
        checkList=[i for i in CurrenyNames if i in findingDoller]
        
        if checkList:
            loopvalues.append('1')    
     fileid=body["value"]
     fileid=fileid.split(",")
     print "Before checking if condition ",loopvalues
     if len(loopvalues) > 0:
        print "inside if cond-------------------"
        emptylist=[]
        removedotlist=[]
        multiplefiles=[]
        name=[]
        listofdata=[]
        withoutdot=[]
        newlist=[]
        CompleteArray = []
        print "-----old code---------"
        for i in fileid:
            file_contents = TextListDetails.objects.get(text_list_id__id=i).file_content_plain_text
            fname = TextList.objects.get(id=i)
            pound=u'\xa3'
            import copy
            orignalfileData=copy.deepcopy(file_contents)
            file_contents=file_contents.replace(pound,"$")
            filesplit=file_contents.split(" ")
            values=re.sub("\$([1-9][0-9,]+)", "",file_contents)
            currencyNames=[]
            newlist=[]
            onlySpecificTerm=' '.join(body["leftandright"].split(' ')[0:2])
            Names=["dollar","dollars","pound","pounds","amount of","USD","SGD","EUR","euros","CNY","CAN","NZD","JPY","yen","RAND","CHF","franc","TWD","NT","AUD","SEK","Krona","rupee","rupees","contract","contract of","value","value of","values at","valued at","valuing at","valuation","valuation of","range","range of","ranges from","ranged from","ranging from","rent","rent of","rental","rental of"]
            modiflyvalues=values.split(" ")
            for i, j in enumerate(filesplit):
                if j in Names:
                   print "modiler values==>",filesplit[i-1]
                   currencyNames.append(filesplit[i-1])
            print "currency values before amount of",currencyNames
            listOfNames=["contract","contract of","value","value of","values at","valued at","valuing at","valuation","valuation of","range","range of","ranges from","ranged from","ranging from","rent","rent of","rental","rental of"]
            for terms in listOfNames:
                findData = [ m.end(0) for m in re.finditer(terms, orignalfileData)]
                for i in findData:
                    data=orignalfileData[i+1]
                    data = ''
                    for j in orignalfileData[i+1:]:
       
                        if j is not ' ' and ord(j) !=32:
                       
                           data = data + j
                        else:
                           break;    
                    currencyNames.append(data)
                    print "currency data======>",data
                print "currency values after amount of",currencyNames
            newCurrencyData=[]
            for i in currencyNames:
                newCurrencyData.append(re.sub("[^0-9]", ".",i))

            loopvalues=[i for i in filesplit if i not in modiflyvalues]
            emptylist=[]
            removedotlist=[]
            removeUnicoe=[x.replace('$','').encode('UTF8') for x in loopvalues]
            emptylist.append(removeUnicoe)
            emptylist=[i.replace('$','') for i in emptylist[0]]
            print "emptylist===",emptylist
            removechar=[]
            for listdata in emptylist:
                
                removechar.append(re.sub("[^0-9]", ".",listdata))

            print "removechar data===",removechar
            for i in newCurrencyData:
                 removechar.append(i)
            filename=fname.name.split(",")
            name.append(filename)
            
            multiplefiles.append([i.replace('...','') for i in removechar])
            
            CompleteArray = []
            for i in multiplefiles:
                currentArray = []
                for j in i:
                   if (re.search(r'\d',j)):
                      currentArray.append(j)
                CompleteArray.append(currentArray)
            print "CompleteArray=====>",CompleteArray

            
            
     else:
        print "------------in else part------"
        CompleteArray=[]
        listofdata=[]
        name=[]
        for i in fileid:
            fname = TextList.objects.get(id=i)
            filename=fname.name.split(",")
            name.append(filename)
        print "=====000====name",name
        listofdata=[]
        for addmultifiles in range(len(name)):
            CompleteArray.append(str(0))
     import copy
     temp=copy.deepcopy(name)
     for i in name:
         listofdata.append(i)
     for i in range(len(CompleteArray)):
         listofdata[i].extend(CompleteArray[i])
     df = pd.DataFrame(listofdata)
     df.transpose()
     filename=[]
     df.to_csv(base_url +'range.csv', index=False)
     os.system("mv "+base_url +"range.csv "+path_url)
     
     msg={"msg":"Document updated succssfully","succss":"true","diff":CompleteArray,'file':temp}
     return JsonResponse(msg)
         


@csrf_exempt
def createFolder(request, project_id, user_id):
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)   
    print body['name'],project_id,user_id
    result = ConceptListFolder(name = body['name'],user_id_id = user_id,project_id_id = project_id )
    result.save()
    msg = {"msg" : "Folder created successfully", "success":True}
    return JsonResponse(msg)

@csrf_exempt    
def getConceptFolder(request, project_id):
    data = ConceptListFolder.objects.filter(project_id = project_id)
    jsondata = serializers.serialize("json", data)
    print (jsondata)
    return HttpResponse(jsondata,status=200,content_type='application/json')

@csrf_exempt
def renameConceptFolder(request, folder_id):
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    print body
    textobj=ConceptListFolder.objects.get(id=folder_id)
    textobj.name=body["name"]
    textobj.save()
    msg={"msg":"rename project succssfully","succss":"true"}
    return JsonResponse(msg)

@csrf_exempt
def deleteConceptFolder(request,folder_id):
    textobj=ConceptListFolder.objects.filter(id=folder_id).delete()   
    msg={"msg":"delete folder succssfully","succss":"true"}
    return JsonResponse(msg)

@csrf_exempt
def addsubFolder(request,user_id,project_id,):
    print "======projectid",project_id
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    try:
        if 'folder_id' in body:
            folder_id=body['folder_id']
            if folder_id !="/":
                folder_id=int(folder_id.split("/")[0].strip())
                parent_folder = Folder.objects.get(id=folder_id) 
            else:
                    parent_folder = None
        else:
            parent_folder = None

        print "3333333333", body
        folder = Folder(status=1)
        folder.name = body['folder_name']
        folder.project_id = project_id
        folder.created_by = folder.updated_by = UserInfo.objects.get(id=user_id)
        folder.status = 1
        folder.save()

        if parent_folder:
            folder.parent_folder_id = parent_folder
            folder.save()
        else:
            print "project id=",project_id
            # map with folder type
            foldertype_obj = FolderType(folder=folder, project_id=project_id,type="text-list")
            foldertype_obj.save()
            print folder, folder.id
        msg={"msg":"Folder created successfully","success":"true"}
        return JsonResponse(msg)
    except Exception,ex:
           print ex

@csrf_exempt
def getprojectData(request, project_id):
    data = Folder.objects.all()

@csrf_exempt
def addUrl(request, project_id):
    body_unicode = request.body.decode("utf-8")
    body = json.loads(body_unicode)
    result = Url(name = body['name'],url = body['url'], folder_id_id=body['folder_id'], project_id_id = project_id )
    result.save() 
    msg = {"msg" : "url created successfully", "success":True}
    return JsonResponse(msg)

@csrf_exempt
def deleteUrl(request, url_id):
    textobj=Url.objects.filter(id=url_id).delete()   
    msg={"msg":"delete Url succssfully","succss":"true"}
    return JsonResponse(msg)


@csrf_exempt
def getUrls(request, project_id):
    data = Url.objects.filter(project_id = project_id)[::-1]
    jsondata = serializers.serialize("json", data)
    return HttpResponse(jsondata,status=200,content_type='application/json')

@csrf_exempt
def conceptStatus(request,user_id):
    userobj=UserInfo.objects.get(userid = user_id)
    userobj.concept_status = not userobj.concept_status
    userobj.save()
    msg={"msg":"User Concept succssfully Changed","user_id":user_id,"succss":"true"}
    return JsonResponse(msg)

@csrf_exempt
def stopListStatus(request,user_id):
    userobj=UserInfo.objects.get(userid = user_id)
    userobj.stoplist_status = not userobj.stoplist_status
    userobj.save()
    msg={"msg":"User stoplist status succssfully Changed","user_id":user_id,"succss":"true"}
    return JsonResponse(msg)

@csrf_exempt
def dataComparision(request):
    body_unicode = request.body.decode("utf-8")
    data = json.loads(body_unicode)

    concept_id = data['concept_id']
    allUmbrell = ConceptListValue.objects.filter(concept_list=concept_id).values('id','umbrella')
    umbrellaListData = []
    umbrellaListIds = []
    for all_u in allUmbrell:
        umbrellaListData.append(all_u['umbrella'])
        umbrellaListIds.append(all_u['id']) 
   
    umbrellaName = ConceptList.objects.get(id=concept_id)
    
    ref_urls_id, ref_data_id = data['ref_urls_id'], data['ref_data_id']
    cmp_urls_id, cmp_data_id = data['cmp_urls_id'], data['cmp_data_id']

    resultCount = {} 
    #get ref_urls data
    if len(ref_urls_id)>0:
        resultCount['RefereceUrls'] = getUrlData(ref_urls_id, concept_id, allUmbrell)

    #get cmp_urls data
    if len(cmp_urls_id)>0:
        resultCount['ComparisionUrls'] = getUrlData(cmp_urls_id, concept_id, allUmbrell)

    #get ref_urls data
    if len(ref_data_id)>0:
        resultCount['RefereceData'] = getDocData(ref_data_id, concept_id, allUmbrell)

    #get cmp_urls data
    if len(cmp_data_id)>0:
        resultCount['ComparisionData'] = getDocData(cmp_data_id, concept_id, allUmbrell)

    resultCount['Umbrella'] = umbrellaListData
    resultCount['Umbrella_name'] = umbrellaName.name
    resultCount['UmbrellaIds'] = umbrellaListIds
    print resultCount
    return JsonResponse(resultCount)

@csrf_exempt
def getFrequentCount(all_data, allUmbrell, docId):
    umbrellData = []
    umbrellData.append(docId)
    zeroUmbrellaCountData = []
    tempFile = all_data.lower()
    tempUmbrellaTermName = []
    for data in allUmbrell:
        tId, umbrellaName = data['id'], data['umbrella'].encode('ascii','ignore')
        tempUmbrellaTermName.append(umbrellaName.lower())
        termData = ConceptListUmbrellaValues.objects.filter(concept_list_value=tId).values('term')
        umbrellaCount = {}
        countData =0
        listContent = [] #
        for term in termData:
            termName =  term['term'].encode('ascii','ignore').lower()
            if not termName in listContent: #
                if not termName in tempUmbrellaTermName: #
                    tempUmbrellaTermName.append(termName)
                findData = [(m.start(0), m.end(0)) for m in re.finditer(termName, tempFile)] #start from here
                for data in findData:
                    frontData = data[0] - 1 #
                    if not ((ord(tempFile[data[1]]) >= 97 and ord(tempFile[data[1]]) <= 122) or (ord(tempFile[frontData]) >= 97 and ord(tempFile[frontData]) <=122)):
                        countData = countData + 1
            listContent.append(termName)#
        umbrellData.append( countData )
        # if countData > 0:
        #     print umbrellaName,'------------',countData
        #     umbrellData.append( countData )
        # else:
        #     zeroUmbrellaCountData.append( countData )
    # umbrellData.extend(zeroUmbrellaCountData)
    return umbrellData
    # print '---------------------', umbrellData,'--------------'
    # return json.dumps({ 'results': umbrellData})

import copy
@csrf_exempt
def getUrlData(urlIds, concept_id,allUmbrell):
    resultCount = {}
    all_data = ''
    for ids in list(set(urlIds)):
        try:
            data = Url.objects.get(id = ids)
            urlData = urllib.urlopen(data.url).read()
            soup = BeautifulSoup(urlData)
            resultCount[data.name.encode('ascii','ignore')] = getFrequentCount(soup.get_text(), allUmbrell, ids)
        except Exception,ex:
            print "Exception"
    return resultCount

@csrf_exempt
def getDocData(doc_id,concept_id,allUmbrell):
    resultCount = {}
    data = ''
    for ids in list(set(doc_id)):
        try:
            # data = data + (TextListDetails.objects.get(text_list_id=ids).file_content_plain_text + " ").encode('ascii','ignore')
            text_list_data = TextListDetails.objects.get(text_list_id=ids)
            text_data = text_list_data.file_content_plain_text.encode('ascii','ignore')
            resultCount[text_list_data.text_list_id.name.encode('ascii','ignore')] = getFrequentCount(text_data, allUmbrell, ids)
            # allData.append(data)
        except Exception,ex:
            print "Exception"
    return resultCount

@csrf_exempt
def getSelectedUrlData(request):
    body_unicode = request.body.decode("utf-8")
    data = json.loads(body_unicode)
    terms = ConceptListUmbrellaValues.objects.filter(concept_list_value = data['umbrella_id']).values_list('term')
    terms = [d[0].encode('ascii','ignore') for d in terms]    
    umbrellaName = ConceptListValue.objects.get(id = data['umbrella_id'])
    terms.append(umbrellaName.umbrella.encode('ascii','ignore'))
    temp_umbrella = list(set([ tmp.lower() for tmp in terms ]))
    text_data = ''
    try:
        url_row_data = Url.objects.get(id = data['url_id'])
        urlData = urllib.urlopen(url_row_data.url).read()
        soup = BeautifulSoup(urlData)
        text_data = soup.get_text()
    except Exception, ex:
        print "Exception"    
    return JsonResponse({'data': seperateData(text_data, temp_umbrella)})

@csrf_exempt
def seperateData(data, temp_umbrella ):
    findData = ''
    tempData = copy.deepcopy(data)
    data = data.lower()
    dataCount = []
    rightCount = {}
    for temp in temp_umbrella:
        findData = [(m.start(0), m.end(0)) for m in re.finditer(temp, data)]
        for position in findData:
            frontData = data[position[0]-1]
            if not ((ord(data[position[1]]) >= 97 and ord(data[position[1]]) <= 122) or (ord(frontData) >= 97 and ord(frontData) <=122)):
                leftstop = 1
                rightstop = 1
                lpos = copy.deepcopy(position[0])
                rpos = copy.deepcopy(position[1])
                count = 0
                middleword = tempData[position[0] : position[1]]
                while leftstop:
                    if lpos == 0 or count == 25:
                        leftstop = 0
                        break
                    if data[lpos] == ' ' and count <= 25 :
                        count += 1
                    lpos -= 1
                leftword  = tempData[lpos+1:position[0]]
                count = 0
                while rightstop:
                    if rpos == len(data) or count == 25:
                        rightstop = 0
                        break
                    if data[rpos] == ' ' and count <= 25:
                        count += 1
                    rpos += 1                
                rightword = tempData[position[1]:rpos]
                dataCount.append({'leftword': leftword,'middleword':middleword,'rightword':rightword })
    return dataCount

@csrf_exempt
def getDocumentData(request):
    body_unicode = request.body.decode("utf-8")
    data = json.loads(body_unicode)    
    terms = ConceptListUmbrellaValues.objects.filter(concept_list_value = data['umbrella_id']).values_list('term')
    terms = [d[0].encode('ascii','ignore') for d in terms]
    umbrellaName = ConceptListValue.objects.get(id = data['umbrella_id'])
    terms.append(umbrellaName.umbrella.encode('ascii','ignore'))
    temp_umbrella = list(set([ tmp.lower() for tmp in terms ]))
    text_data = ''
    try:
        text_data = TextListDetails.objects.get(text_list_id__id=data['doc_id']).file_content_plain_text
        text_data = text_data.replace("%",'')         
    except Exception, ex:
        print "Exception"    
    return JsonResponse({'data': seperateData(text_data, temp_umbrella)})

def CombinationHead(request,left,right,conceptlist):
    oneLine=' '.join(left.split(' ')[-15:]) + ' '.join(right.split(' ')[:15])
    emptylist=[]
    for i in conceptlist:
        valuessplit=i.split(',')
        for j in valuessplit:
            for newvalue in  oneLine.split(" "):
                if j == newvalue:
                   emptylist.append(j)
    try:
        termName=[i for i in emptylist if i][0]
        termId=ConceptListUmbrellaValues.objects.filter(term=termName)[0]
        termHead=ConceptListUmbrellaValues.objects.filter(concept_list_value_id=termId.concept_list_value_id)[0]
        return termHead.term.upper()        
    except IndexError:
           return None
def convert_pdf_to_txt(path):
    rsrcmgr = PDFResourceManager()
    retstr = StringIO()
    codec = 'utf-8'
    laparams = LAParams()
    device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
    fp = file(path, 'rb')
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    password = ""
    maxpages = 0
    caching = True
    pagenos=set()
    for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages, password=password,caching=caching, check_extractable=True):
        interpreter.process_page(page)
    fp.close()
    device.close()
    str = retstr.getvalue()
    retstr.close()
    return str


@csrf_exempt    
def UserManaul(request):
    import shutil
    print "-------in user manula-----",request.FILES['user_manual']
    myfile = request.FILES['user_manual']
    directory=settings.MEDIA_URL + str(uuid.uuid4()) + "/"
    fs = FileSystemStorage(location=directory)
    filename = fs.save(myfile.name, myfile)
    uploaded_file_url = fs.url(directory+filename)
    print "filename====>",filename
    print "filename====>",uploaded_file_url
    pdfData=document_to_text(request,filename,directory+filename)
    subprocess.call(['soffice', '--headless', '--convert-to', 'pdf', directory+filename],cwd=base_url)
    if os.path.exists(path_url+filename):
       os.unlink(path_url+filename)
    shutil.move(base_url+filename,path_url)
    return JsonResponse({"msg":"SuccessFully uploaded pdf","pdfname":filename})


def getNewConceptLists(request,project_id):
    project = Projects.objects.get(id=project_id)
    try:
        lst_concept_lists = []
        lst_concept_list_obj = ConceptList.objects.filter(project_id=project)
        for x in lst_concept_list_obj:
            dict_concept_list = {}
            lst_concept_list_values = []
            lst_concept_list_value_obj = ConceptListValue.objects.filter(concept_list=x)
            lst_umbrella = []
            for y in lst_concept_list_value_obj:

                dict_umbrella = {}
                dict_umbrella.update({"umbrella": y.umbrella})
                    # lst_umbrella.append(y.umbrella)
                lst_concept_list_umbrella_values_obj = ConceptListUmbrellaValues.objects.filter(concept_list_value=y)
                for z in lst_concept_list_umbrella_values_obj:
                    lst_terms_obj = ConceptListUmbrellaValues.objects.filter(
                        concept_list_value=z.concept_list_value)
                    terms = ""

                    lst_terms=[]
                    for term_obj in lst_terms_obj:

                        terms = terms + term_obj.term + ","
                        lst_terms.append(term_obj.term)
                    terms = terms[:-1]
                    dict_umbrella.update({"terms": terms, "id": z.id,"lst_terms":lst_terms})
                lst_umbrella.append(dict_umbrella)
            dict_concept_list.update({"name": x.name, "values": lst_umbrella, "id": x.id})
            lst_concept_lists.append(dict_concept_list)

    except Exception, ex23:
           print  ex23

    #print "concept_lists=====", json.dumps(lst_concept_lists)
    return lst_concept_lists

